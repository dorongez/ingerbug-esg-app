import streamlit as st
from openai import OpenAI
import os
import io
import docx
import json
from PyPDF2 import PdfReader
import pandas as pd
import base64
from collections import defaultdict
from fpdf import FPDF
import requests

# Initialize OpenAI client
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

st.set_page_config(page_title="GingerBug ESG Assistant", layout="wide")
st.title("🌱 GingerBug - Release your sustainable power")

if st.button("🔄 Start Over"):
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    st.rerun()

st.markdown("### ✉️ Enter your email to personalize your experience")
user_email = st.text_input("Your email address (we'll use this to save your session or send you your report)", "")
if user_email:
    st.session_state.user_email = user_email

st.markdown("### 🏢 Company Info")
st.session_state.company_name = st.text_input("Company Name", st.session_state.get("company_name", ""))
st.session_state.company_url = st.text_input("Company Website URL", st.session_state.get("company_url", ""))
st.session_state.country = st.text_input("Country", st.session_state.get("country", ""))
st.session_state.report_goal = st.multiselect("📊 Reporting Goals", ["EcoVadis", "VSME", "CSRD Prep", "GRI"], default=["EcoVadis"])

if st.session_state.get("company_url"):
    try:
        st.image(f"https://logo.clearbit.com/{st.session_state.company_url}", width=100)
    except:
        st.warning("Could not load logo from Clearbit")

lang = st.selectbox("🌐 Choose language", ["English", "Français", "Deutsch", "Español"])

for key in ['summaries', 'drafts', 'generated_metrics', 'checklist_progress']:
    if key not in st.session_state:
        st.session_state[key] = [] if key == 'summaries' else {}

uploaded_files = st.file_uploader("📄 Upload ESG documents (PDF, DOCX, XLSX)", accept_multiple_files=True)

st.markdown("### ✅ Prefilled Checklist by Reporting Goal")
goal_to_checklist = {
    "EcoVadis": ["Code of Conduct", "Supplier Policy", "Diversity & Inclusion", "Environmental Policy"],
    "CSRD Prep": ["Double Materiality Assessment", "Energy Audit", "GRI Alignment", "Sustainability Governance Policy"]
}
for goal in st.session_state.report_goal:
    checklist = goal_to_checklist.get(goal, [])
    if checklist:
        st.markdown(f"**{goal} requirements:**")
        for item in checklist:
            key = f"progress_{goal}_{item}"
            all_docs = [doc['file'] for doc in st.session_state.summaries] + list(st.session_state.get("drafts", {}).keys())
            is_checked = item in all_docs
            st.session_state.checklist_progress[key] = st.checkbox(item, value=is_checked, key=key)

def extract_text(file):
    if file.name.endswith(".pdf"):
        reader = PdfReader(file)
        return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file.name.endswith(".xlsx") or file.name.endswith(".xls"):
        df = pd.read_excel(file)
        return df.to_string()
    else:
        return file.read().decode(errors="ignore")

def summarize_text(text):
    prompt = f"Summarize this ESG-related document content:\n{text[:3000]}"
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3
    )
    return response.choices[0].message.content.strip()

if uploaded_files:
    st.markdown("## 📝 Uploaded File Summaries")
    for file in uploaded_files:
        content = extract_text(file)
        summary = summarize_text(content)
        key_safe = file.name.replace("/", "_").replace(".", "_")
        st.session_state.summaries.append({"file": file.name, "summary": summary})
        st.text_area(f"Summary - {file.name}", summary, height=150, key=f"summary_{key_safe}_{len(st.session_state.summaries)}")

if st.session_state.summaries:
    completeness = len(st.session_state.summaries) / 5
    st.markdown("### 🔦 ESG Readiness")
    if completeness >= 1:
        st.success("🟢 Ready: You have uploaded all key documents.")
    elif completeness >= 0.6:
        st.warning("🟠 In Progress: You have uploaded most of the documents.")
    else:
        st.error("🔴 Incomplete: Please upload more ESG documents.")

if st.button("✨ Generate Missing Policies"):
    with st.spinner("Generating missing policy drafts..."):
        missing = [item for goal in st.session_state.report_goal for item in goal_to_checklist.get(goal, []) if item not in st.session_state.get("drafts", {})]
        for topic in missing:
            try:
                prompt = f"Create a basic draft of a {topic} for a company in {st.session_state.country} named {st.session_state.company_name}. Please base it only on credible sources."
                response = client.chat.completions.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.3
                )
                st.session_state.drafts[topic] = response.choices[0].message.content.strip()
                st.success(f"✅ {topic} generated.")
            except Exception as e:
                st.error(f"⚠️ Failed to generate {topic}: {str(e)}")

if st.session_state.get("drafts"):
    st.markdown("## 📄 Drafted Policies")
    for title, text in st.session_state.drafts.items():
        key_safe = title.replace(" ", "_").replace("/", "_")
        st.text_area(f"{title}", text, height=200, key=f"draft_{key_safe}_{len(st.session_state.drafts)}")
        doc = docx.Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(text)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button(label=f"📥 Download {title}.docx", data=buffer, file_name=f"{title}.docx")

if st.button("📌 Suggest Metrics & KPIs"):
    with st.spinner("Generating KPIs based on selected frameworks..."):
        goals = ", ".join(st.session_state.get("report_goal", []))
        prompt = f"Suggest 5 ESG metrics and KPIs for a company reporting under {goals}."
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3
        )
        st.session_state.generated_metrics = response.choices[0].message.content.strip()

if st.session_state.generated_metrics:
    st.markdown("### 📈 Recommended ESG Metrics and KPIs")
    st.text_area("Suggested KPIs", st.session_state.generated_metrics, height=250)

class PDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_font("Arial", size=12)

    def safe_text(self, text):
        return text.encode('latin-1', 'replace').decode('latin-1')

if st.button("🗕️ Export Full ESG Report"):
    with st.spinner("Bundling full ESG report..."):
        pdf = PDF()
        pdf.add_page()

        pdf.multi_cell(0, 10, pdf.safe_text("🔍 ESG Report Summary"))
        pdf.multi_cell(0, 10, pdf.safe_text(f"Company: {st.session_state.get('company_name', '')}\nEmail: {user_email}\nCountry: {st.session_state.get('country', '')}\nGoals: {', '.join(st.session_state.get('report_goal', []))}\n"))

        pdf.add_page()
        pdf.multi_cell(0, 10, pdf.safe_text("✅ Checklist Completion:"))
        for key, value in st.session_state.checklist_progress.items():
            status = "✅" if value else "❌"
            pdf.multi_cell(0, 10, pdf.safe_text(f"{status} {key.split('_')[-1]}"))

        if st.session_state.generated_metrics:
            pdf.add_page()
            pdf.multi_cell(0, 10, pdf.safe_text("📈 Suggested ESG KPIs"))
            pdf.multi_cell(0, 10, pdf.safe_text(st.session_state.generated_metrics))

        if st.session_state.drafts:
            pdf.add_page()
            pdf.multi_cell(0, 10, pdf.safe_text("📄 Generated Policies"))
            for topic, content in st.session_state.drafts.items():
                pdf.multi_cell(0, 10, pdf.safe_text(f"{topic}\n{content}\n"))

        if st.session_state.summaries:
            pdf.add_page()
            pdf.multi_cell(0, 10, pdf.safe_text("📂 Uploaded Document Summaries"))
            for entry in st.session_state.summaries:
                pdf.multi_cell(0, 10, pdf.safe_text(f"{entry['file']}\n{entry['summary']}\n"))

        full_pdf_name = f"GingerBug_Full_Report_{st.session_state.get('company_name', 'report')}.pdf"
        pdf.output(full_pdf_name)
        with open(full_pdf_name, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
            href = f'<a href="data:application/pdf;base64,{b64}" download="{full_pdf_name}">🗕️ Download Full ESG PDF</a>'
            st.markdown(href, unsafe_allow_html=True)

st.markdown("### 📊 ESG Report Summary Dashboard")
if st.session_state.summaries:
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Uploaded Documents", len(st.session_state.summaries))
        st.metric("Generated Policies", len(st.session_state.drafts))
        total_items = sum(1 for k in st.session_state.checklist_progress)
        completed = sum(1 for k, v in st.session_state.checklist_progress.items() if v)
        st.metric("Checklist Completion", f"{completed}/{total_items}")
    with col2:
        st.markdown("**Reporting Goals:**")
        for goal in st.session_state.report_goal:
            st.markdown(f"- ✅ {goal}")

    st.markdown("---")
    st.markdown("**Next Steps:**")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        if st.button("➕ Upload More Docs"):
            st.warning("Scroll up to use the uploader.")
    with col2:
        if st.button("📝 Generate Policies"):
            st.rerun()
    with col3:
        if st.button("📊 Refine KPIs"):
            st.rerun()
    with col4:
        if st.button("🗕 Download PDF"):
            st.rerun()
